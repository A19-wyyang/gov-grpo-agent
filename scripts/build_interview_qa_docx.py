from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/阿里百炼算法岗_简历模拟面试问答_杨文宇.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)


def set_run_font(run, name="Microsoft YaHei", size=11, color=BLACK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, *, size, color=BLACK, bold=False, before=0, after=6, line=1.25) -> None:
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def add_para(doc, text="", *, size=11, color=BLACK, bold=False, italic=False,
             before=0, after=6, line=1.25, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text: str, level=0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_callout(doc, label: str, text: str, *, fill=CALLOUT, label_color=DARK_BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.20
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, size=10.5, color=label_color, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5)


def add_qa(doc, number: int, question: str, answer: str, probes: str | None = None,
           boundary: str | None = None) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"Q{number}. {question}")
    set_run_font(r, size=12, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r1 = p.add_run("回答示例：")
    set_run_font(r1, size=10.5, color=NAVY, bold=True)
    r2 = p.add_run(answer)
    set_run_font(r2, size=10.5)
    if probes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.20
        r1 = p.add_run("继续追问：")
        set_run_font(r1, size=10, color=CAUTION, bold=True)
        r2 = p.add_run(probes)
        set_run_font(r2, size=10, color=MUTED)
    if boundary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.20
        r1 = p.add_run("回答边界：")
        set_run_font(r1, size=10, color=RED, bold=True)
        r2 = p.add_run(boundary)
        set_run_font(r2, size=10, color=MUTED)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.80)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    configure_style(styles["Normal"], size=10.5, after=5, line=1.25)
    configure_style(styles["Heading 1"], size=16, color=BLUE, bold=True, before=16, after=8, line=1.10)
    configure_style(styles["Heading 2"], size=13, color=BLUE, bold=True, before=12, after=6, line=1.10)
    configure_style(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True, before=8, after=4, line=1.10)
    for list_style in ("List Bullet", "List Bullet 2", "List Number"):
        configure_style(styles[list_style], size=10.5, after=4, line=1.25)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("阿里百炼算法岗日常实习 | 简历模拟面试问答")
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])
    return doc


def add_cover(doc: Document) -> None:
    add_para(doc, "", after=40)
    add_para(doc, "INTERVIEW PREP GUIDE", size=10, color=BLUE, bold=True,
             after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "阿里百炼算法岗日常实习", size=27, color=NAVY, bold=True,
             after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "简历模拟面试问答手册", size=21, color=DARK_BLUE, bold=True,
             after=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "AI 搜索 / Agent 方向", size=14, color=MUTED,
             after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc,
        "定位",
        "根据杨文宇简历与阿里巴巴 ATH MaaS 业务线招聘信息整理。用于面试前逐题演练，不建议逐字背诵。",
        fill=LIGHT_BLUE,
    )
    add_para(doc, "", after=88)
    add_para(doc, "候选人：杨文宇", size=11, color=NAVY, bold=True,
             after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "准备方向：Agentic RL / RAG / Deep Search / Multi-Agent / SFT",
             size=10.5, color=MUTED, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "版本：2026 年 6 月", size=10, color=MUTED,
             after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_overview(doc: Document) -> None:
    doc.add_heading("使用说明", level=1)
    add_callout(
        doc,
        "面试策略",
        "回答采用“结论先行 -> 自己负责的模块 -> 一处关键细节 -> 指标或复盘 -> 边界说明”的顺序。遇到没真正做过的方向，不要补造经历；先承认边界，再给出可落地方案。",
        fill=LIGHT_BLUE,
    )
    add_para(doc, "岗位匹配判断", size=12, color=DARK_BLUE, bold=True, before=4, after=5)
    add_bullet(doc, "强匹配：Agentic RL、轨迹采样、工具调用、Reward Design、SFT、模型评估。")
    add_bullet(doc, "可迁移：将政务 Agent 的 runtime、Verifier、Judge、监控体系迁移到 Deep Search 与 Multi-Agent。")
    add_bullet(doc, "需要补强：经典信息检索指标、召回与重排、Query Rewrite、搜索增强推理、Agentic Memory。")
    add_bullet(doc, "面试风险：简历写得较强，面试官会连续追问实现细节、指标定义、数据泄漏和真实贡献边界。")

    doc.add_heading("招聘信息拆解", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1900, 3550, 3910])
    headers = ["岗位关键词", "面试可能考察", "你的连接点"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    rows = [
        ("Deep Search / RAG", "召回、重排、摘要、评估、延迟", "补齐检索基础；用 Verifier 思路讲结果可信度"),
        ("Query Rewrite", "意图识别、多轮上下文、改写收益", "可从诉求识别与槽位补全迁移"),
        ("Multi-Agent", "角色拆分、路由、协作、失败恢复", "可从工具链编排与结构化动作迁移"),
        ("Agentic RL", "rollout、reward、GRPO、坍塌、hacking", "简历第一项目，是你的主战场"),
        ("模型训练评估", "SFT、LoRA、数据构造、指标、幻觉", "简历第二项目，强调小样本受控扩增"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.4)
    add_para(doc, "", after=2)

    doc.add_heading("推荐演练顺序", level=2)
    add_number(doc, "先练 Q1-Q5：保证开场、自我介绍、岗位动机和项目概述稳定。")
    add_number(doc, "重点练 Q6-Q24：第一项目是面试官最容易问穿的位置。")
    add_number(doc, "再练 Q25-Q36：把已有 Agent 经验迁移到搜索、RAG 和 Multi-Agent。")
    add_number(doc, "最后练 Q37-Q51：第二项目、基础知识、行为题和反问。")
    add_callout(doc, "事实边界", "本文答案基于简历和本地教学项目。涉及真实线上规模、具体模型超参数、业务转化指标、线上 A/B 数值时，请只回答你能确认的内容。", fill="FFF8E8", label_color=CAUTION)
    doc.add_page_break()


def add_section(doc, title: str, intro: str, qas: list[tuple]) -> None:
    doc.add_heading(title, level=1)
    add_para(doc, intro, size=10.5, color=MUTED, after=8)
    for qa in qas:
        add_qa(doc, *qa)


def main() -> None:
    doc = setup_document()
    add_cover(doc)
    add_overview(doc)

    q = 1
    section_1 = [
        (q, "请做一个 90 秒自我介绍。", "您好，我叫杨文宇，目前是西安电子科技大学通信工程硕士。我关注大模型应用算法，重点是 Agent 策略优化和 SFT。最近一段经历中，我参与了政务办理 Agent 的策略优化，把诉求识别、信息追问、政策查询、资格核验、材料检查和提交或拒答建模为多轮决策任务。我主要负责 case schema、结构化动作、trajectory log、Verifier + Judge 奖励，以及策略坍塌和 reward hacking 排查。另一段项目中，我围绕 Qwen3-8B 做劳动仲裁文书 LoRA SFT，完成小样本受控扩增、assistant-only loss 和离线评估。我希望把这些经验迁移到百炼的 AI 搜索和 Agent 场景，继续做可评估、可诊断、能落地的系统。", "如果只给 30 秒，保留：背景、两个项目关键词、为什么匹配岗位。", "不要主动报手机、邮箱；不要声称已做过 Deep Search 线上落地。"),
        (q := q + 1, "为什么投阿里百炼 ATH MaaS 业务线，而不是泛大模型岗位？", "我看重的是岗位把 AI 搜索和 Agent 放在真实业务链路里，而不是只做单点模型指标。招聘信息里有 Deep Search、RAG、Query Rewrite、排序摘要、Multi-Agent、Agentic Reasoning、Memory 和 RL，这些模块都需要端到端评估和故障诊断。我的第一段经历正好训练了我从 runtime、轨迹、奖励到监控闭环看问题的方式。我的检索工程经验还需要继续补强，但 Agentic RL 和评估诊断能力可以直接迁移。", "为什么你认为 Agent 经验能迁移到搜索？", "表达为“可迁移 + 待补强”，不要假装已经熟悉百炼内部架构。"),
        (q := q + 1, "你认为自己的核心优势是什么？", "第一，我能把模糊业务流程转成可训练、可评估的结构化任务；第二，我对 reward 设计保持警惕，会同时看平均 reward、组内方差、动作分布和工具调用率，不会只盯一个数字；第三，我做过小样本 SFT 数据构造，知道模型效果不只是调参问题，数据边界和评估协议同样重要。", "举一个你发现指标误导的例子。", None),
        (q := q + 1, "你最需要补强的地方是什么？", "我需要补强的是搜索系统的系统性工程经验，尤其是召回、重排、Query Rewrite 和 Deep Search 评估。我的准备方式不是泛泛看概念，而是把它们映射到已有 Agent 框架：明确状态、动作、工具观测、失败标签和离线指标，再补实现。我能快速进入这个方向，但不会把待学习内容说成已经完成。", "如果入职第一周让你做 Query Rewrite，你怎么开始？", None),
        (q := q + 1, "你的通信工程背景如何帮助做 Agent 算法？", "通信工程训练了我对系统链路、信号噪声、指标分解和约束优化的敏感度。Agent 系统也不是只看一个生成结果，而是要观察状态、动作、工具返回、错误传播和最终收益。专业不同不是优势本身，真正有价值的是我能把系统思维转成工程诊断方法。", None, None),
    ]
    add_section(doc, "第一部分：开场与岗位匹配", "目标：开场稳定、事实准确、主动把话题引到你最熟悉的 Agentic RL 项目。", section_1)

    section_2 = [
        (q := q + 1, "用两分钟介绍政务办理 Agent 策略优化项目。", "这个项目面向线上政务事项办理。难点是用户信息往往不完整，Agent 需要多轮追问，并按政策约束依次做资格、材料和风险核验。我们把任务建模为多轮决策过程：case 中拆分用户诉求、待补全槽位、政策规则、材料状态和标准结果；动作空间包括 Ask_User、Policy_Search、Eligibility_Check、Material_Check、Risk_Check、Submit 和 Refuse。Agent 与环境交互生成 trajectory，每一步统一记录 state、action、observation。奖励侧将规则 Verifier 和表达 Judge 聚合成 trajectory-level reward，同一 case 采样多条 rollout 后用组内相对优势做 GRPO 训练。我的重点工作是环境 schema、奖励信号、坍塌监控和 reward hacking 修复。", "让你画架构图时，按 case -> runtime -> rollout -> Verifier/Judge -> group advantage -> training -> monitoring 顺序画。", None),
        (q := q + 1, "为什么这是多轮决策任务，而不是普通问答？", "因为最终回复是否好，不只取决于语言质量，还取决于流程是否完整。比如用户申请社保补贴，资格满足但材料缺失，如果 Agent 没调用 Material_Check 就直接给出泛化建议，看起来通顺但业务上不可靠。多轮建模可以让每一步动作都可观测、可回放、可打分。", "状态、动作、观测分别是什么？", None),
        (q := q + 1, "你的 case schema 具体有哪些字段？", "我会分成四层。第一层是 visible，包含用户初始诉求和已知槽位；第二层是 hidden_truth，由环境持有，用于模拟用户追问结果、材料状态和风险标记；第三层是 policy_rules，通过政策查询工具返回，包括必填槽位、必要工具、材料要求和资格阈值；第四层是 expected_result，仅供离线评估器使用。关键隔离原则是策略不能读取 expected_result，避免答案泄漏。", "为什么 policy_rules 不直接全部暴露给 Agent？", "如果面试官问真实系统字段，按你真实经历补充；这里给出的是结构化回答框架。"),
        (q := q + 1, "动作空间为什么用结构化动作？", "结构化动作有三个价值：第一，环境可以确定性执行和返回工具观测；第二，Verifier 能检查必要工具、槽位和顺序；第三，日志更容易聚合分析。生产环境可以让模型输出 JSON action，再由 runtime 做 schema 校验。自由文本仍然保留在最终回复或工具参数中，但核心决策必须结构化。", "模型输出非法 JSON 怎么处理？", None),
        (q := q + 1, "trajectory log 里记录什么？", "每一步至少记录 state snapshot、action、observation、slot_status、tool_history、failure_tags 和最终 decision。训练需要 reward，排障更需要可回放性。比如发现材料缺失场景表现异常时，可以沿轨迹确认模型是否调用了 Material_Check、工具返回了什么、Judge 为什么仍给高分。", "为什么只保存最终回复不够？", None),
        (q := q + 1, "为什么选择 GRPO？", "同一个政务 case 可以自然采样多条办理路径，适合组内比较。GRPO 使用同组 rollout 的相对 reward 构造 advantage，不需要额外训练 value model。对 Agent 来说，这样可以比较完整但稍长的正确路径、过早提交路径、过度拒答路径和话术规避路径。它比只做单条监督学习更适合优化过程策略。", "与 PPO、DPO 相比有什么区别？", None),
        (q := q + 1, "GRPO 的 advantage 怎么计算？", "简化表达是：同一 case 的每条 rollout 先得到 trajectory reward，再计算组内均值和标准差，advantage_i 等于 reward_i 减组内均值，再除以组内标准差。高于同组平均的路径得到正 advantage，低于平均的路径得到负 advantage。真实训练中还会结合策略比率、clip、KL 约束等目标更新 token-level policy。", "如果组内标准差接近 0 怎么办？", None),
        (q := q + 1, "Verifier 与 Judge 如何分工？", "Verifier 负责硬事实：槽位是否补齐、资格是否核验、材料检查是否调用、风险判断是否完成、最终动作是否与标准结果一致。Judge 只读取脱敏摘要和最终回复，评价表达完整性、清晰度和可执行性。硬事实不能交给 Judge，否则模型会利用语言偏好绕过流程。", "Judge 是否会看到完整 trajectory？", None),
        (q := q + 1, "trajectory-level reward 如何聚合？", "可以表达为 reward = verifier_weight * hard_score + judge_weight * quality_score + process_score - penalties。penalty 单独处理 early_submit、missing_required_tool、repeated_question、over_refuse、wrong_final_action 等错误。具体权重需要在离线集上做消融和分场景审计，不应该凭直觉只调一个数字。", "如果规则过多，会不会限制探索？", None),
        (q := q + 1, "为什么环境要允许模型犯错，而不是直接拦截错误 Submit？", "训练环境需要让错误动作发生并留下标签，才能形成负样本。如果环境层直接禁止所有错误动作，策略看不到错误路径，也无法学到为什么它低分。生产 runtime 可以分层处理：高风险动作增加硬门禁，但离线训练和回放环境仍保留错误轨迹用于学习与诊断。", "训练环境与线上环境如何保持一致？", None),
        (q := q + 1, "你如何发现策略坍塌？", "我们不只看平均 reward，还看 action 分布、Refuse 和 Submit 占比、必要工具调用率、policy entropy 和同组 reward 方差。异常现象是 rollout 越来越集中到 Refuse 或过早 Submit，同一组里的 trajectory 差异缩小，GRPO 对比信号变弱。定位后发现错误提交惩罚偏强，而合理探索奖励不足。", "如果平均 reward 上升但熵下降，你如何判断？", None),
        (q := q + 1, "entropy bonus 的作用是什么？", "entropy bonus 鼓励策略保留一定多样性，减少过早收敛到固定动作路径。但它不是 reward 修复的替代品。如果奖励目标本身错了，只加 entropy 仍然是在错误目标附近探索。因此需要同时调整拒答惩罚、工具调用 reward 和硬事实约束。", "系数怎么定？", "不要给出未经确认的线上系数。可回答通过离线 sweep、场景分桶和稳定性指标选择。"),
        (q := q + 1, "请具体讲 reward hacking 案例。", "在资格基本满足但材料缺失的 case 中，Agent 有时跳过 Material_Check，直接输出“建议补齐相关材料后再办理”这类泛化话术。因为表达完整、语气清晰，Judge 会给较高分，但流程其实不可靠。我们回溯 trajectory 和 Judge 理由后，确认 Verifier 没有惩罚必要工具缺失。修复方式是增加 missing-tool penalty，要求 Submit 前完成资格、材料和风险检查，同时降低 Judge 权重。", "如果最终动作是 Refuse，是否也要求调用所有工具？", None),
        (q := q + 1, "missing-tool penalty 会不会导致无意义工具滥用？", "会有这个风险，所以不能简单要求所有 case 调用所有工具。更合理的是根据事项类型、当前状态和终止动作定义条件化必需工具集合，并对重复调用、无效调用和超长路径增加成本。目标是最小充分流程，而不是工具调用越多越好。", "如何定义最小充分流程？", None),
        (q := q + 1, "如何防止 expected_result 泄漏给策略？", "从数据接口上隔离。runtime 只向策略提供 visible 信息、用户追问返回和工具观测；expected_result 只进入离线 scorer。还要检查 prompt 拼接、日志回灌、缓存和训练样本构造，避免标准答案通过隐藏字段或模板泄漏。评估集按 case 维度切分，不能让同一案件变体跨训练和验证。", "如何自动化检测泄漏？", None),
        (q := q + 1, "如何评估 Agent，不只看最终成功率？", "我会分四层：结果层看正确提交、正确拒答和人工升级率；过程层看必要工具调用率、槽位补全率、平均轮次和重复动作率；安全层看错误提交率、风险漏检率和政策越权；训练层看组内 reward 方差、entropy 和动作分布。再按事项类型、缺失信息类型和风险等级做切片。", "哪个指标最重要？", None),
        (q := q + 1, "如果线上突然出现 Refuse 比例升高，你怎么排查？", "先确认是流量分布变化还是策略变化。按事项、用户信息完整度和风险等级切片；再看工具可用性、超时率、schema 校验错误和模型版本。随后抽样回放 trajectory，对比 Refuse 前最后几个动作和观测。若是 reward 或训练问题，再看离线组内方差、entropy 和拒答路径优势是否异常。", "如果是工具超时导致呢？", None),
        (q := q + 1, "真实大模型训练与本地教学项目有什么差距？", "本地项目完整实现了 case、runtime、trajectory、reward、组内 advantage 和故障监控，但训练模拟器更新的是行为模式概率，不是 8B 模型参数。真实 GRPO 需要模型按当前策略采样 token 序列，保存旧策略 log probability，计算组内 advantage，再用 clipped objective、KL 约束和 entropy bonus 做 GPU 反向传播。这个边界必须讲清。", "你在真实项目中负责到了哪一层？", "按真实参与程度作答。不要把本地教学仓库说成公司线上代码。"),
        (q := q + 1, "如果让你把该系统扩展到线上，你会增加哪些工程机制？", "我会增加四类机制：第一，动作 schema 校验和高风险终止动作门禁；第二，工具超时、重试、降级和幂等；第三，prompt、模型、工具和 policy rule 的版本化；第四，线上回放与分桶监控，包括错误提交、风险漏检、工具调用率、延迟和 token 成本。灰度阶段还要保留人工复核。", "如何做灰度发布？", None),
    ]
    add_section(doc, "第二部分：第一项目深挖 - 政务办理 Agentic RL", "目标：这是主项目。回答必须经得住“为什么、怎么做、如何验证、哪里可能错”的连续追问。", section_2)

    section_3 = [
        (q := q + 1, "如果把你的 Agentic RL 经验迁移到 Deep Search，你会如何拆系统？", "我会拆成 Query Understanding、Query Rewrite、检索规划、搜索工具调用、结果去重与重排、证据聚合、答案生成、引用校验和评估监控。Agent runtime 负责状态和动作，搜索工具返回 observation，Verifier 检查引用覆盖、证据一致性和必要搜索步骤，Judge 评价答案组织和可读性。训练时可以对同一查询采样不同搜索路径做组内比较。", "哪些模块适合规则，哪些适合模型？", "这是方案题，不要说成已在线上完成。"),
        (q := q + 1, "RAG 的基本链路是什么？", "典型链路是文档预处理、切分、向量或关键词索引、查询理解、召回、重排、上下文组装、生成和引用校验。实际系统通常不是只用向量检索，而是结合 BM25、dense retrieval、metadata filter 和 reranker。评估要拆开看检索质量与生成质量，避免只看最终回答。", "chunk size 如何选择？", None),
        (q := q + 1, "召回和重排分别解决什么问题？", "召回阶段追求高 recall，用较低成本从大规模库中拿到候选；重排阶段追求 precision 和顺序质量，用更强模型对较小候选集排序。召回漏掉的文档，重排救不回来；重排不足则会让上下文被噪声占用。两者要用分层指标分别评估。", "常见召回指标有哪些？", None),
        (q := q + 1, "Query Rewrite 为什么重要？", "用户查询可能短、歧义、多轮依赖强或口语化。Rewrite 可以补充实体、时间、约束和上下文，也可以拆成子查询提升召回。但改写不能擅自改变意图，所以需要保留原查询、做约束校验，并评估改写前后的 recall、NDCG、最终回答准确率和延迟。", "什么时候不该 rewrite？", None),
        (q := q + 1, "Deep Search 与普通 RAG 的区别是什么？", "普通 RAG 通常是一次召回后生成；Deep Search 更强调多步规划、迭代搜索、证据缺口识别、来源比较和带引用综合。它需要判断“当前证据是否足够”，必要时发起新的子查询。代价是延迟和成本上升，所以要有停止条件和预算控制。", "停止条件怎么设计？", None),
        (q := q + 1, "如何评估 Deep Search？", "我会分为检索、证据、答案和效率四层。检索看 Recall@K、MRR、NDCG；证据看引用覆盖率、来源质量和证据是否支持结论；答案看事实正确性、完整性和可读性；效率看搜索轮次、延迟、token 和工具成本。复杂问题还要按是否需要多跳、是否时效敏感做切片。", "Judge 能否直接评所有维度？", None),
        (q := q + 1, "如何设计 Multi-Agent System？", "先确认是否真的需要多 Agent。如果任务可以由单 Agent 加工具完成，优先保持简单。需要拆分时，可以按职责划分 planner、searcher、evidence reviewer、writer 和 verifier，明确每个 Agent 的输入、输出 schema、最大轮次和失败处理。共享状态要结构化，避免多个 Agent 用长文本互相污染上下文。", "多 Agent 常见失败模式是什么？", None),
        (q := q + 1, "Agentic Memory 怎么做？", "我会区分短期工作记忆、会话记忆和长期用户记忆。短期记忆服务当前任务，保留计划、已调用工具和证据；会话记忆帮助多轮上下文；长期记忆必须经过用户授权、敏感信息过滤、时效管理和可删除机制。检索时要做相关性、时间和权限过滤，不能简单把历史全部塞入 prompt。", "如何避免错误记忆长期污染？", None),
        (q := q + 1, "搜索型 Agent 如何防 prompt injection？", "把网页和文档内容当作不可信数据，而不是系统指令。工具返回需要做来源隔离、内容标注和解析；模型 prompt 中明确禁止执行检索内容里的指令；敏感工具调用增加 allowlist、参数校验和人工确认；最终引用要能回溯到原始来源。还要构造注入攻击集做离线评估。", "如果网页要求忽略之前指令呢？", None),
        (q := q + 1, "如何控制搜索 Agent 的延迟和成本？", "使用预算化策略：限制最大搜索轮次、每轮候选数、上下文 token 和高成本模型调用；先用轻量模型做 query 分类与路由，难题再升级；缓存稳定检索结果；并行执行独立子查询；对证据充分的简单问题提前停止。评估时同时看质量、P95 延迟和单请求成本。", "什么时候值得用更强模型？", None),
        (q := q + 1, "如果搜索结果互相矛盾，Agent 应该怎么办？", "先识别冲突维度，例如时间、地区、来源权威性和适用条件。保留多个来源并显式说明差异，优先选择更权威、更新、更贴合问题约束的证据。如果仍不能确定，应表达不确定性或请求补充信息，而不是强行合成一个结论。", "如何自动判断来源权威性？", None),
        (q := q + 1, "长上下文模型是否会替代 RAG？", "不会完全替代。长上下文减少了部分切分和召回压力，但仍有成本、时效、权限、噪声和证据定位问题。RAG 的价值不只是节省窗口，还包括动态更新、可引用、权限过滤和检索可解释性。更合理的是长上下文与检索协同。", None, None),
    ]
    add_section(doc, "第三部分：岗位迁移题 - AI 搜索、RAG 与 Multi-Agent", "目标：这些题不是要求你假装做过，而是考察能否把已有能力迁移到岗位方向。", section_3)

    section_4 = [
        (q := q + 1, "介绍 Qwen3-8B 仲裁文书 SFT 项目。", "项目面向劳动仲裁文书生成，输入案件事实、证据、仲裁请求、主体和金额等结构化材料，输出仲裁查明、仲裁认定和裁决结果。企业真实样本约 300 条，直接 LoRA SFT 容易过拟合模板。我主要做受控数据扩增，把案件拆成完整生成、分段生成、主体或金额一致性纠错、信息不足拒答等任务，扩展为约 2000 条 ChatML 样本；训练使用 assistant-only loss；评估使用 Base 与 SFT 同集同参数对比。", "为什么不是直接让大模型自由生成扩增样本？", None),
        (q := q + 1, "300 条扩展到 2000 条，如何控制质量？", "核心是受控任务拆解，而不是无约束改写。每条样本继承真实案件 JSON 中的事实字段，只改变训练任务形式，例如完整文书、分段生成、纠错和拒答。需要校验主体、金额、请求和证据字段是否一致，并按案件维度切分训练与验证，避免同一案件的不同任务变体跨集合造成泄漏。", "如何抽检？", None),
        (q := q + 1, "什么是 assistant-only loss？", "ChatML 样本包含 system、user 和 assistant。监督目标是 assistant 侧的文书输出，因此通过 mask 让 system 和 user token 不参与 loss，只对 assistant token 计算交叉熵。否则模型会把输入材料也当成生成目标，造成训练目标偏移，并浪费梯度容量。", "mask 如何实现？", None),
        (q := q + 1, "为什么使用 LoRA？", "8B 模型全量微调显存和训练成本较高，而业务数据规模有限。LoRA 冻结基座参数，在部分线性层增加低秩可训练矩阵，以较少参数适配任务。优势是成本低、迭代快、便于保留多个业务 adapter；代价是表达能力受 rank 和注入位置影响，需要通过验证集选择。", "rank、alpha、target modules 怎么选？", "如未确认真实训练参数，回答选择方法，不要现场编造具体数字。"),
        (q := q + 1, "为什么不能只用 BLEU 或 ROUGE 评估法律文书？", "法律文书存在结构、主体、金额、请求响应和事实一致性要求。文本相似不代表事实正确，合理措辞变化也可能被相似度误伤。因此我们重点看结构完整率、主体一致率和幻觉率，并在相同验证集、prompt 和解码参数下对比 Base 与 SFT。", "幻觉率怎么定义？", None),
        (q := q + 1, "你如何降低幻觉？", "第一，数据层确保输入和目标字段对齐，并加入信息不足拒答；第二，任务层加入主体和金额一致性纠错；第三，推理层限制解码随机性并明确只依据材料；第四，评估层把无依据新增事实单独统计。若上线，还可以增加规则校验和引用回溯。", "如果规则校验和生成冲突怎么办？", None),
        (q := q + 1, "如何避免模板过拟合？", "让训练任务和输出形态更丰富，同时保持事实受控。除了完整文书，增加分段生成、字段纠错和拒答；验证集按案件切分；观察模型是否机械复制固定段落、是否在信息不足时仍生成确定性结论。必要时降低训练轮次、调整 LoRA 容量或扩充覆盖更多案件类型。", None, None),
        (q := q + 1, "离线指标提升是否能代表上线效果？", "不能完全代表。离线集用于快速比较版本，但线上还要观察真实输入缺失、格式噪声、长尾案件、用户编辑率、人工采纳率和错误严重度。上线前应做人工盲评和小流量灰度，并保留回退。", None, None),
    ]
    add_section(doc, "第四部分：第二项目深挖 - Qwen3-8B LoRA SFT", "目标：强调数据质量、训练 mask 和评估协议。不要把小样本扩增讲成简单复制。", section_4)

    section_5 = [
        (q := q + 1, "请解释 Transformer self-attention。", "对输入 hidden states 分别线性映射得到 Q、K、V。使用 Q 与 K 的点积衡量 token 间相关性，除以 sqrt(d_k) 控制数值尺度，经过 softmax 得到权重，再对 V 加权求和。多头注意力允许不同子空间关注不同关系。训练时可并行，但自回归生成仍逐 token 解码。", "为什么除以 sqrt(d_k)？", None),
        (q := q + 1, "交叉熵损失是什么？", "对正确 token 的预测概率取负对数并求平均。语言模型学习的是给定上下文时下一个 token 的条件概率。assistant-only loss 本质上是在交叉熵外增加 mask，只对目标回复 token 聚合损失。", "label shift 怎么做？", None),
        (q := q + 1, "temperature、top-p 对生成有什么影响？", "temperature 调整 logits 的平滑程度，越低越确定；top-p 从累计概率达到阈值的最小 token 集合采样。法律文书和政务办理更重事实稳定性，通常偏保守；需要 rollout 多样性时可以提高采样随机性，但要结合奖励和安全约束。", None, None),
        (q := q + 1, "BM25 与向量检索有什么区别？", "BM25 基于词项匹配和逆文档频率，对精确关键词、专有名词和稀有实体很有效；向量检索强调语义相似，适合表达不同但含义接近的查询。生产系统常做 hybrid retrieval，再由 reranker 统一排序。", "什么时候 BM25 反而更好？", None),
        (q := q + 1, "Recall@K、MRR、NDCG 分别是什么？", "Recall@K 看前 K 个结果是否覆盖相关文档；MRR 关注第一个相关结果的排名倒数，适合答案通常对应单个关键结果的场景；NDCG 考虑多档相关性和排序位置，适合多个结果都有价值的场景。指标要根据业务目标选择。", "如果 Recall 高但答案质量差，排查什么？", None),
        (q := q + 1, "PPO、GRPO、DPO 怎么区分？", "PPO 是在线策略优化，通常依赖 advantage 估计和 value model；GRPO 利用同一问题多条采样的组内相对 reward 构造 advantage，减少对 value model 的依赖；DPO 通常使用偏好对直接优化策略，不需要在线 rollout reward。选择取决于数据形态、成本和任务是否需要过程探索。", None, None),
        (q := q + 1, "模型评估为什么需要切片？", "整体均值会掩盖长尾问题。Agent 要按事项类型、信息缺失类型、工具失败、风险等级和对话轮次切片；搜索要按查询难度、多跳需求、时效敏感和领域切片；文书生成要按案件类型、字段缺失和文本长度切片。切片才能定位可修复的具体问题。", None, None),
        (q := q + 1, "工具调用 Agent 出现循环调用怎么办？", "runtime 层设置最大轮次、重复动作检测和幂等键；策略层对无效重复调用增加 penalty；观测层让工具返回结构化状态，帮助模型识别已经完成的步骤；评估层统计重复动作率和超限率。必要时触发人工升级。", None, None),
    ]
    add_section(doc, "第五部分：算法与工程基础", "目标：回答短而准确。不会时先讲核心定义，再说明工程含义。", section_5)

    section_6 = [
        (q := q + 1, "讲一个你排查问题的完整过程。", "我会选 reward hacking 案例。先看到材料缺失场景的 Material_Check 调用率下降，但部分回复的 Judge 分仍较高；然后按 case 回放 trajectory，发现模型跳过材料检查，输出泛化建议；再查看 Judge 理由和 Verifier 规则，确认硬事实约束缺口；最后增加 missing-tool penalty、降低 Judge 权重，并观察材料检查调用率和组内 reward 方差恢复。这个过程说明平均分不足以诊断 Agent。", "如果修复后延迟上升怎么办？", None),
        (q := q + 1, "如果你和同事对 reward 权重有分歧，怎么解决？", "先把争议转成可验证假设。准备场景分桶和失败案例，固定模型版本，分别跑权重组合，比较结果正确率、风险漏检、工具调用率、平均轮次和组内方差。高风险错误设置硬约束，表达质量再做软权重优化。用实验协议而不是偏好争论。", None, None),
        (q := q + 1, "你如何快速学习一个陌生方向？", "先画最小链路，再建立可运行基线和可观察指标。比如补 Deep Search，我会先实现 query -> hybrid retrieval -> rerank -> evidence -> answer 的基线，再增加 rewrite 和多步搜索，每次只引入一个变量。并构造小型错误集验证 recall、引用和停止条件，而不是先堆复杂架构。", None, None),
        (q := q + 1, "你的简历里最容易被质疑的地方是什么？", "第一项目容易被追问真实训练边界，所以我会明确区分 runtime、奖励、监控与 token-level GPU 训练；第二项目容易被追问 300 到 2000 条是否造成泄漏，所以要讲案件级切分和受控扩增；岗位方向还会追问搜索基础，我会诚实说明待补强并给出迁移方案。", None, None),
        (q := q + 1, "实习时间和投入如何保证？", "我会根据课程安排给出明确、可兑现的每周到岗天数和连续实习周期，并提前说明考试或学校节点。对日常实习而言，稳定投入比模糊承诺重要。", None, "请根据你的真实可用时间补充具体数字。"),
        (q := q + 1, "你有什么问题想问面试官？", "我会问三个问题：第一，团队当前 AI 搜索与 Agent 的核心业务场景是什么，最关注质量、延迟还是成本中的哪类指标？第二，实习生通常会负责哪一段链路，是检索、评估、Agent runtime 还是训练优化？第三，团队目前最想解决的失败模式是什么，例如多步搜索稳定性、引用可信度、工具调用还是长尾评估？", "避免只问福利和转正；优先问业务问题和成功标准。", None),
    ]
    add_section(doc, "第六部分：行为题、压力题与反问", "目标：展示事实意识、排障方法和学习能力。不要把压力题回答成防御姿态。", section_6)

    doc.add_page_break()
    doc.add_heading("临场速记卡", level=1)
    add_callout(doc, "主线", "业务流程结构化 -> rollout 轨迹 -> Verifier / Judge -> GRPO 组内优势 -> 监控坍塌 -> 修复 reward hacking。", fill=LIGHT_BLUE)
    doc.add_heading("必须主动说清的边界", level=2)
    add_bullet(doc, "策略不能读取 `expected_result`；标准答案只进入离线评估。")
    add_bullet(doc, "Judge 只评表达，硬事实交给 Verifier。")
    add_bullet(doc, "entropy bonus 不能替代 reward 修复。")
    add_bullet(doc, "本地教学仓库不是 8B token-level GPU 训练。")
    add_bullet(doc, "搜索方向是可迁移能力，不要包装成已完成线上 Deep Search 项目。")
    doc.add_heading("白板架构", level=2)
    add_para(doc, "Case / Query -> Runtime -> Structured Action -> Tool Observation -> Trajectory -> Verifier + Judge -> Reward -> Same-case Rollout Group -> Advantage -> Policy Update -> Monitoring", size=10.5, color=NAVY, bold=True, after=7)
    doc.add_heading("重点指标", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2600, 3300, 3460])
    for idx, text in enumerate(("场景", "重点指标", "防止误判")):
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
        r = table.cell(0, idx).paragraphs[0].add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for row in [
        ("Agentic RL", "正确提交率、风险漏检、工具调用率、熵、组内方差", "不要只看平均 reward"),
        ("Deep Search", "Recall@K、NDCG、引用覆盖、事实正确、P95 延迟", "拆开检索与生成"),
        ("文书 SFT", "结构完整、主体金额一致、幻觉率、人工采纳", "同集同参数对比，案件级切分"),
    ]:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.5)
    add_para(doc, "", after=4)
    add_heading = doc.add_heading
    add_heading("最后检查", level=2)
    add_bullet(doc, "每个回答先给结论，再给一处实现细节。")
    add_bullet(doc, "被追问数字时，只说能确认的指标；不能确认就讲评估方法。")
    add_bullet(doc, "不确定时先澄清场景，不要急着堆术语。")
    add_bullet(doc, "把面试当成一次系统设计和故障复盘讨论。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "阿里百炼算法岗日常实习 - 简历模拟面试问答手册"
    doc.core_properties.subject = "AI 搜索 / Agent 方向面试准备"
    doc.core_properties.author = "杨文宇"
    doc.core_properties.keywords = "Agentic RL, GRPO, RAG, Deep Search, Multi-Agent, SFT"
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
