from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "阿里算法岗日常实习_简历定制面试题与参考回答_杨文宇.docx"

BLUE = "1F4D78"
ACCENT = "2E74B5"
DARK = "0B2545"
MUTED = "64748B"
INK = "1F2937"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WARN = "FFF4CC"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, bold=None, italic=None, color=INK, east_asia="Microsoft YaHei", latin="Calibri"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths):
            set_cell_width(row.cells[index], width)
            set_cell_margins(row.cells[index])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_border_left(paragraph, color=ACCENT, size="14", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def set_keep(paragraph, next_paragraph=False, together=False):
    p_pr = paragraph._p.get_or_add_pPr()
    if next_paragraph:
        p_pr.append(OxmlElement("w:keepNext"))
    if together:
        p_pr.append(OxmlElement("w:keepLines"))


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_font(style, size, color=INK, bold=False, latin="Calibri", east_asia="Microsoft YaHei"):
    style.font.name = latin
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold


def setup_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    style_font(normal, 11, INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, BLUE, 10, 5),
    ):
        style = styles[name]
        style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    question = styles.add_style("Question", 1)
    style_font(question, 11.2, DARK, True)
    question.paragraph_format.space_before = Pt(10)
    question.paragraph_format.space_after = Pt(3)
    question.paragraph_format.line_spacing = 1.12

    answer = styles.add_style("Answer", 1)
    style_font(answer, 10.5, INK)
    answer.paragraph_format.space_after = Pt(5)
    answer.paragraph_format.line_spacing = 1.18

    note = styles.add_style("Note", 1)
    style_font(note, 9.5, MUTED)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.15

    code = styles.add_style("Code Block", 1)
    style_font(code, 9, DARK, latin="Consolas", east_asia="Microsoft YaHei")
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.line_spacing = 1.05


def add_numbering(doc, bullet=True):
    numbering = doc.part.numbering_part._element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_bullet(doc, text, bullet_num_id, bold_prefix=None):
    p = doc.add_paragraph()
    apply_numbering(p, bullet_num_id)
    set_paragraph_spacing(p, after=4, line=1.18)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, size=10.3, bold=True, color=DARK)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run, size=10.3)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.3)
    return p


def add_numbered(doc, text, decimal_num_id):
    p = doc.add_paragraph()
    apply_numbering(p, decimal_num_id)
    set_paragraph_spacing(p, after=4, line=1.18)
    set_run_font(p.add_run(text), size=10.3)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, after=0, line=1.0)
    run = p.add_run("阿里算法岗日常实习｜简历定制面试准备")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [7000, 2360], indent=0)
    set_table_borders(table, color=WHITE, size="0")
    set_repeat_table_header(table.rows[0])
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    p_left = left.paragraphs[0]
    p_right = right.paragraphs[0]
    set_paragraph_spacing(p_left, after=0, line=1.0)
    set_paragraph_spacing(p_right, after=0, line=1.0)
    set_run_font(p_left.add_run("候选人：杨文宇｜Agent 算法 / 应用算法"), size=8.5, color=MUTED)
    add_page_number(p_right)


def add_title(doc, text, size=24, color=DARK, after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, after=after, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return p


def add_subtitle(doc, text, size=13, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, after=after, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def add_callout(doc, label, body, fill=CALLOUT):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8, line=1.15)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    paragraph_shading(p, fill)
    paragraph_border_left(p)
    run = p.add_run(f"{label} ")
    set_run_font(run, size=10.2, bold=True, color=DARK)
    run = p.add_run(body)
    set_run_font(run, size=10.2, color=INK)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5, add_spacer=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(value), size=font_size, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.05)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index > 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(str(value)), size=font_size, color=INK)
    if add_spacer:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2)
    return table


def add_qa(doc, index, question, answer, tips=None, formula=None, fill_item=None):
    p = doc.add_paragraph(style="Question")
    set_keep(p, next_paragraph=True, together=True)
    set_run_font(p.add_run(f"Q{index}. {question}"), size=11.2, bold=True, color=DARK)

    label = doc.add_paragraph()
    set_keep(label, next_paragraph=True, together=True)
    set_paragraph_spacing(label, after=1, line=1.0)
    set_run_font(label.add_run("参考回答"), size=9.3, bold=True, color=ACCENT)

    for block in answer:
        p = doc.add_paragraph(style="Answer")
        set_run_font(p.add_run(block), size=10.5, color=INK)
    if formula:
        p = doc.add_paragraph(style="Code Block")
        paragraph_shading(p, LIGHT_GRAY)
        set_run_font(p.add_run(formula), size=9, color=DARK, latin="Consolas")
    if tips:
        p = doc.add_paragraph(style="Note")
        paragraph_shading(p, CALLOUT)
        paragraph_border_left(p, color="94A3B8", size="8", space="5")
        set_run_font(p.add_run("面试官可能继续追问："), size=9.3, bold=True, color=BLUE)
        set_run_font(p.add_run(tips), size=9.3, color=MUTED)
    if fill_item:
        p = doc.add_paragraph(style="Note")
        paragraph_shading(p, WARN)
        paragraph_border_left(p, color="D69E2E", size="8", space="5")
        set_run_font(p.add_run("面试前补齐："), size=9.3, bold=True, color="7A5A00")
        set_run_font(p.add_run(fill_item), size=9.3, color="7A5A00")


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_terminal_marker(doc):
    # WPS emits an extra blank page when the document ends directly after a table.
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(1)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = p.add_run("\u200b")
    set_run_font(run, size=1, color=WHITE)


def add_cover(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=8)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=10, line=1.0)
    set_run_font(p.add_run("ALIBABA ALGORITHM INTERNSHIP"), size=10, bold=True, color=ACCENT)
    add_title(doc, "阿里算法岗日常实习", size=25, color=DARK, after=2)
    add_title(doc, "简历定制面试题与参考回答", size=21, color=DARK, after=12)
    add_subtitle(doc, "Agent 算法 / 应用算法方向｜候选人：杨文宇", size=13, color=BLUE, after=4)
    add_subtitle(doc, "基于 2026 年 5 月 30 日版本简历整理", size=10.5, color=MUTED, after=22)
    add_callout(
        doc,
        "使用方式",
        "先用 90 秒版本建立主线，再重点复习 Agentic RL 与 SFT 两段项目。文中黄色提示为简历未披露但面试中高概率被追问的数据，请只填写真实记录。",
        fill=LIGHT_BLUE,
    )
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=40, after=5)
    set_run_font(p.add_run("内容范围"), size=10, bold=True, color=MUTED)
    for line in (
        "开场表达与岗位匹配",
        "政务办理 Agent：GRPO、Verifier / Judge、策略坍塌、Reward Hacking",
        "仲裁文书生成：Qwen3-8B、LoRA SFT、assistant-only loss、离线评估",
        "算法基础、场景设计与行为面试",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        set_paragraph_spacing(p, after=3, line=1.0)
        set_run_font(p.add_run(line), size=10.5, color=INK)


def add_overview(doc, bullet_num_id):
    doc.add_heading("一、面试官视角：这份简历会怎样被追问", level=1)
    add_callout(
        doc,
        "一句话判断",
        "你的简历不是“调包跑模型”型项目，而是围绕 Agent 决策质量做训练环境、奖励信号、失败模式定位和修复。面试成败取决于能否把工程现象解释成可验证的算法问题。",
    )
    add_table(
        doc,
        ["项目", "面试官最关心的问题", "你需要证明什么"],
        [
            ("政务 Agent GRPO", "动作空间、reward 可验证性、坍塌与 hacking", "你真正参与了训练闭环，而不是只写业务流程"),
            ("仲裁文书 SFT", "300→2000 数据构造、assistant-only loss、指标定义", "你能控制小样本过拟合并做公平评估"),
            ("两段项目共性", "如何把业务约束转成训练与评估信号", "你具备应用算法岗需要的闭环能力"),
        ],
        [1800, 3900, 3660],
        font_size=9.4,
    )
    doc.add_heading("面试表达原则", level=2)
    for text in (
        "先给结论，再讲证据链：现象 → 统计指标 → 根因假设 → 改动 → 验证结果。",
        "主动区分“简历已有事实”和“面试前需补齐的数据”。没有真实记录时，不要临场虚构百分比。",
        "对于政务流程，强调硬约束优先于语言表达质量；对于法律文书，强调结构、主体、金额和幻觉控制优先于文本相似度。",
        "对无法确定的问题，给出可验证的实验设计，比强行给结论更专业。",
    ):
        add_bullet(doc, text, bullet_num_id)


def add_opening(doc):
    doc.add_heading("二、开场与岗位匹配", level=1)
    add_qa(
        doc,
        1,
        "请用 90 秒做一个自我介绍。",
        [
            "面试官您好，我叫杨文宇，目前是西安电子科技大学通信工程硕士在读。本科和硕士阶段的专业训练让我比较习惯从系统约束、信号反馈和实验验证三个角度拆问题。",
            "最近两段项目都围绕大模型应用算法展开。第一段是在政务办理 Agent 场景中做策略优化：我参与构建多轮工具调用环境，设计 Verifier + Judge 奖励，并基于 GRPO 做训练；训练过程中重点排查了策略过早集中到 Refuse 或 Submit，以及跳过必要工具调用获取高表达分的问题。第二段是基于 Qwen3-8B 的仲裁文书 LoRA SFT：我将约 300 条案件级样本扩成约 2000 条受控多任务样本，使用 assistant-only loss，并搭建 Base / SFT 同集对比评估。",
            "我希望做应用算法或 Agent 算法方向的实习，因为我比较擅长把业务流程拆成可训练、可评估、可迭代的算法闭环，也希望在更复杂的真实业务中提升工程化能力。",
        ],
        tips="你在项目中最核心的个人贡献是什么？为什么通信工程背景适合做 Agent 算法？",
    )
    add_qa(
        doc,
        2,
        "为什么想做阿里的应用算法或 Agent 算法实习？",
        [
            "我希望做的是有真实业务约束的 Agent，而不是只在开放问答上优化表达。阿里的业务场景通常包含搜索、推荐、客服、交易、办公或企业服务等复杂链路，Agent 要在效率、准确性、合规和用户体验之间做平衡，这和我在政务办理项目中处理硬事实校验、必要工具调用和最终表达质量的思路是一致的。",
            "我能带来的价值是：先把流程结构化，再建立 trajectory 级日志和评估信号，最后根据失败模式迭代策略。对实习阶段而言，我也希望进一步补齐大规模训练、在线评估和工程部署经验。",
        ],
        tips="如果岗位更偏搜索、推荐或客服，你的经验如何迁移？",
    )
    add_qa(
        doc,
        3,
        "两段项目里，哪一段最能体现你的算法能力？",
        [
            "政务 Agent 项目更能体现我处理闭环问题的能力。因为它不是只优化最终回复，而是把诉求识别、槽位追问、政策查询、资格核验、材料检查、风险判断和提交或拒答建模为多轮决策任务。",
            "真正困难的部分是奖励和失败模式诊断：训练早期出现动作分布向 Refuse 或过早 Submit 集中，我通过 action 分布、组内 reward 方差和必要工具调用率定位问题；后来又发现模型利用 Judge 的表达偏好绕过 Material_Check，我回溯 trajectory log 后加入 missing-tool penalty 并降低 Judge 权重。这种从现象到验证再到修复的过程，是我认为最有代表性的部分。",
        ],
    )


def add_agent_rl(doc):
    doc.add_heading("三、政务办理 Agent：GRPO 策略优化", level=1)
    add_callout(
        doc,
        "回答主线",
        "这段经历最重要的不是背 GRPO 公式，而是说明：如何把流程变成环境，如何让 reward 可信，如何用日志诊断策略问题，以及修复后怎样证明行为真的改善。",
    )
    add_qa(
        doc,
        4,
        "请从业务到算法完整介绍政务办理 Agent 项目。",
        [
            "业务上，线上政务事项办理不是一次性问答，而是一个带约束的多轮流程。用户可能缺少信息，政策规则需要查询，资格、材料和风险条件需要核验，最后才能提交或拒答。",
            "算法上，我将其建模为多轮决策任务。每个 case 包含用户诉求、待补全槽位、政策规则、材料状态和标准结果；Agent 通过结构化动作和环境交互生成 trajectory。训练侧使用 Verifier 校验硬事实，Judge 评估最终表达，再将 trajectory-level reward 用于同一 case 多条 rollout 的组内对比和 GRPO 更新。",
            "我重点负责环境和日志结构、奖励信号设计，以及策略坍塌和 reward hacking 的定位与修复。核心目标是让模型不仅会说，还能按必要流程正确做事。",
        ],
    )
    add_qa(
        doc,
        5,
        "你如何把政务办理流程建模成 Agent 环境？",
        [
            "我把每个事项拆成 case schema：用户诉求、已知与缺失槽位、政策规则、材料状态、风险条件、标准结论。状态记录当前已收集信息、工具返回、已执行动作和剩余约束；动作空间使用结构化 schema，避免自由文本直接驱动环境。",
            "环境根据动作返回 observation，并在 trajectory log 中记录每一步 state、action、observation。这样既能支持训练，也能在失败时回放具体路径，判断是模型理解错、工具调用缺失，还是 reward 设计有漏洞。",
        ],
        formula="Ask_User(slot) | Policy_Search(query) | Eligibility_Check(rule_ids)\nMaterial_Check(material_ids) | Risk_Check(flags) | Submit(result) | Refuse(reason)",
        tips="状态是否包含历史对话全文？工具异常怎么处理？动作参数校验在哪里做？",
    )
    add_qa(
        doc,
        6,
        "为什么动作要结构化？直接让模型自然语言调用工具不行吗？",
        [
            "结构化动作有三个好处。第一，环境可以确定性解析，减少格式漂移。第二，Verifier 能基于动作类型和参数做硬校验，例如 Submit 前是否执行 Material_Check。第三，统计 action 分布、工具调用率和路径差异时更可靠。",
            "自然语言仍然用于用户沟通，但执行层最好有稳定 schema。生产系统中还需要加入参数合法性校验、重试、超时、幂等和工具错误回传。",
        ],
    )
    add_qa(
        doc,
        7,
        "你为什么选择 GRPO？和 SFT、DPO、PPO 相比有什么考虑？",
        [
            "SFT 适合让模型先学会基本流程和动作格式，但难以直接优化长程决策质量。DPO 更适合已有偏好对的数据；在这个项目中，我们能够对同一个 case 采样多条 rollout，并通过 Verifier 和 Judge 对完整轨迹评分，所以更适合使用组内相对比较。",
            "GRPO 不依赖单独训练 value model，而是用同组样本 reward 的相对优势做更新，工程上更简洁。它的前提是同组 rollout 要有足够差异，reward 也要能区分好坏；这正是后续策略坍塌问题需要重点监控的原因。",
        ],
        tips="PPO 的 critic 有什么作用？没有 value model 的代价是什么？",
    )
    add_qa(
        doc,
        8,
        "请解释 GRPO 的组内相对优势。",
        [
            "对同一个 case 采样 G 条轨迹，先得到每条轨迹的 reward。然后用组内均值和标准差标准化，得到相对优势。直觉上，同一题里的优秀轨迹获得正优势，较差轨迹获得负优势，减少了不同 case 难度差异对更新的干扰。",
            "如果同组轨迹几乎都走同一路径，或者 reward 几乎相同，标准化后可用信号会变弱。因此我会监控组内 reward 方差、动作分布熵和路径差异，而不是只看平均 reward。",
        ],
        formula="A_i = (r_i - mean(r_1 ... r_G)) / (std(r_1 ... r_G) + epsilon)",
        fill_item="实际 rollout group size、是否做 reward clipping、是否使用 KL 正则及其系数。",
    )
    add_qa(
        doc,
        9,
        "Verifier 和 Judge 为什么要分层？",
        [
            "两者解决的问题不同。Verifier 负责可确定判断的硬事实，例如槽位是否补全、资格是否核验、材料检查是否执行、Submit 前必要工具是否齐全。它应该尽量规则化、可回放、可解释。",
            "Judge 只读取脱敏摘要和最终回复，评价完整性、清晰度和可执行性。Judge 更适合处理表达质量，但不应该覆盖流程硬约束。项目中出现 reward hacking 的根本原因，就是 Judge 表达偏好权重过高，而 Verifier 对必要工具缺失惩罚不足。",
        ],
        tips="Judge 模型和策略模型是否同源？如何做 judge calibration？如何避免敏感信息进入 Judge？",
    )
    add_qa(
        doc,
        10,
        "你的 trajectory-level reward 如何设计？",
        [
            "我会先将奖励拆成硬约束分、过程分和表达分。硬约束优先级最高，包括关键槽位、资格、材料、风险和最终结果一致性；过程分关注必要工具和合理步骤；表达分由 Judge 评估最终回复。",
            "聚合时不能只追求加权和好看，还要设置门槛和惩罚。例如缺少必要工具时，即使最终话术完整，也不能获得高分；错误 Submit 应有明显惩罚。实际权重需要通过失败样本和消融实验校准。",
        ],
        formula="R = w_hard * R_verifier + w_process * R_tools + w_text * R_judge\n    - penalty_wrong_submit - penalty_missing_tool",
        fill_item="各 reward 项真实权重、惩罚阈值、是否有 hard gate，以及调参前后的对比。",
    )
    add_qa(
        doc,
        11,
        "你如何发现策略坍塌？",
        [
            "训练早期，部分 case 的 rollout 逐渐集中到 Refuse 或过早 Submit。表面上 reward 不一定立刻异常，但同组 trajectory 差异变小，GRPO 的组内对比信号会减弱。",
            "我同时查看三类指标：动作分布，尤其是 Refuse 和 Submit 占比；同一 case 内 reward 方差；必要工具调用率。结合轨迹回放后，定位到错误提交惩罚过强、合理探索奖励不足，模型倾向选择固定的低风险路径。",
        ],
        tips="如何区分合理拒答率上升和策略坍塌？是否按 case 类型分桶？",
        fill_item="坍塌前后 Refuse / Submit 分布、组内 reward 方差、必要工具调用率的真实变化。",
    )
    add_qa(
        doc,
        12,
        "你具体怎样修复策略坍塌？为什么 entropy bonus 有效？",
        [
            "我在训练侧引入 entropy bonus，鼓励策略在早期保留一定探索能力；同时重新校准拒答惩罚和工具调用 reward，避免某个动作因为风险最低而成为固定解。",
            "entropy bonus 不是越大越好。过大会让策略长期随机，降低执行稳定性；过小则无法缓解过早收敛。合理做法是结合训练阶段和验证集指标调节系数，并观察动作熵、路径多样性和任务成功率是否同时改善。",
        ],
        formula="L_total = L_GRPO - beta_entropy * H(pi(.|s))",
        fill_item="entropy 系数、是否使用衰减策略、调整前后的关键指标。",
    )
    add_qa(
        doc,
        13,
        "请解释你遇到的 reward hacking。",
        [
            "在“资格基本满足但材料缺失”的 case 中，Agent 频繁跳过 Material_Check，直接输出泛化建议。由于 Judge 偏好表达完整、语气清晰的回答，这类轨迹拿到了偏高的表达分。",
            "我回溯 trajectory log 和 Judge 评分理由，确认模型利用了奖励漏洞：Verifier 没有惩罚必要工具缺失，Judge 权重又足以掩盖流程错误。修复方式是加入 missing-tool penalty，要求 Submit 前完成资格核验、材料检查和风险判断，并降低 Judge 权重。",
            "这个问题说明，Agent reward 不能只评价最终文本，还要覆盖关键过程。否则模型会学习“看起来合理”的捷径。",
        ],
        tips="如何自动发现新的 reward hacking？能否做 adversarial case generation？",
        fill_item="修复前后 Material_Check 调用率、错误 Submit 率、组内 reward 方差的真实数据。",
    )
    add_qa(
        doc,
        14,
        "missing-tool penalty 应该怎样实现？",
        [
            "可以根据事项类型定义 required-tool graph。Verifier 在 trajectory 结束时检查：如果进入 Submit，是否完成该 case 对应的 Eligibility_Check、Material_Check 和 Risk_Check；如果缺失，则施加惩罚或直接触发 hard gate。",
            "同时要避免一刀切。不同事项的必需工具可能不同，工具失败也不能等同于 Agent 漏调。因此规则需要结合 case schema、工具返回状态和异常码判断。",
        ],
        formula="required_tools(case_type, final_action) ⊆ executed_successful_tools(trajectory)",
    )
    add_qa(
        doc,
        15,
        "如何评估 Agent 训练是否真正有效？",
        [
            "平均 reward 只能作为一个信号，不能作为唯一结论。我会至少分四层评估：任务层看正确提交、合理拒答和整体成功率；过程层看必要工具调用率、漏调率、冗余调用率和平均步数；策略层看动作分布、路径多样性、组内 reward 方差和熵；安全层看错误 Submit、敏感信息泄露和异常工具调用。",
            "此外需要按 case 类型分桶，尤其区分材料缺失、资格不满足、信息缺失和风险命中场景。否则总体平均值可能掩盖高风险子类退化。",
        ],
        fill_item="项目实际使用的验证集规模、case 分桶、核心指标定义与最终结果。",
    )
    add_qa(
        doc,
        16,
        "Judge 本身不稳定，如何保证奖励可信？",
        [
            "第一，限定 Judge 职责，只评估表达质量，不让它替代硬事实校验。第二，固定 rubric，把完整性、清晰度和可执行性拆开评分。第三，抽样做人审校准，关注 Judge 与人工的一致性。第四，对关键风险场景做成对对比，验证 Judge 是否偏爱冗长或泛化话术。",
            "如果 Judge 漂移明显，可以使用多 Judge 投票、置信度阈值、规则兜底和定期校准集回归测试。核心原则是让可规则化的问题尽量规则化。",
        ],
    )
    add_qa(
        doc,
        17,
        "reference path 有什么用？会不会限制 Agent 探索？",
        [
            "reference path 是基于政务流程整理的参考执行路径，主要用于定义合理动作顺序、构造校验规则和分析轨迹偏差。它不应该被当成唯一正确路径，因为有些 case 可以跳过不适用步骤，有些信息已经由用户提供。",
            "更稳妥的做法是把它设计成约束图或允许路径集合：明确哪些步骤必须完成，哪些步骤可选，哪些动作有前置条件。这样既保留探索空间，也能守住流程底线。",
        ],
    )
    add_qa(
        doc,
        18,
        "如果把这个 Agent 上线，你还会补什么？",
        [
            "训练之外，我会补四类能力。第一，线上日志与可观测性：完整 trace、工具耗时、失败码和版本信息。第二，工具治理：参数校验、超时重试、幂等、降级和权限控制。第三，安全与隐私：敏感字段脱敏、最小化传输、审计和高风险人工兜底。第四，线上评估：灰度、分桶指标、回归集和人工抽检。",
            "政务场景里，宁可在高风险路径上请求补充信息或转人工，也不能用流畅表达掩盖事实不确定性。",
        ],
    )
    add_qa(
        doc,
        19,
        "这段项目还有哪些局限？",
        [
            "第一，规则 Verifier 的覆盖度依赖 case schema，长尾事项可能需要持续补规则。第二，Judge 仍然可能存在偏好漂移，需要人工校准和回归集。第三，如果 rollout 多样性不足，GRPO 的相对优势信号会变弱。第四，离线提升不等同于线上体验提升，还需要真实流量中的异常处理、时延和成本评估。",
            "后续我会优先补充难例挖掘、自动化回归、基于失败轨迹的 case 生成，以及更系统的消融实验。",
        ],
    )


def add_sft(doc):
    doc.add_heading("四、仲裁文书生成：Qwen3-8B LoRA SFT", level=1)
    add_qa(
        doc,
        20,
        "请完整介绍仲裁文书生成项目。",
        [
            "项目目标是基于 Qwen3-8B 生成劳动仲裁文书中的仲裁查明、仲裁认定和裁决结果等核心内容。企业只有约 300 条案件级 JSON，直接做 LoRA SFT 容易过拟合固定模板。",
            "我的工作分三部分：第一，围绕案件事实、证据、仲裁请求、主体、金额和标准文书字段做受控任务拆解，将约 300 条案件扩展成约 2000 条 ChatML 多任务样本；第二，使用 assistant-only loss，只对 assistant 侧目标文书计算监督损失；第三，搭建 Base / SFT 同集对比评估，在一致条件下比较结构完整性、主体一致性和幻觉率。",
            "离线评估中，结构完整率由 0.71 提升到 0.96，主体一致率由 0.74 提升到 0.93，幻觉率由 0.26 降到 0.09。",
        ],
    )
    add_qa(
        doc,
        21,
        "为什么选择 Qwen3-8B 和 LoRA？",
        [
            "8B 量级在效果、训练成本和部署成本之间比较平衡，适合企业侧验证。LoRA 只训练低秩适配参数，显存占用和迭代成本更低，也便于保留底座能力。",
            "面试中我会把实际选择依据讲清楚：当时可用算力、最长输入长度、推理成本、中文法律文本表现，以及是否对不同底座做过小规模对比。不能只回答“因为常用”。",
        ],
        fill_item="实际底座版本、是否为 Base 或 Instruct、GPU 配置、LoRA rank / alpha / dropout、学习率、epoch、max sequence length。",
    )
    add_qa(
        doc,
        22,
        "300 条案件如何扩展到约 2000 条样本？这算数据增强吗？",
        [
            "它更准确地说是受控任务拆解，而不是随意改写。每个案件可以派生多种训练任务：完整文书生成、分段生成、主体一致性纠错、金额一致性纠错、信息不足拒答等。所有输出都必须受原始案件 JSON 和标准文书字段约束。",
            "这样做的价值是增加监督信号密度，让模型分别学习结构、局部生成、一致性校验和拒答边界。风险是同一案件派生样本可能造成数据泄漏，所以必须先按案件划分训练、验证、测试集，再在各自集合内部拆任务。",
        ],
        tips="有没有使用模型合成数据？如何做质量抽检？派生任务的比例是多少？",
        fill_item="各任务类型的真实样本量与比例；拆分规则；人工抽检比例。",
    )
    add_qa(
        doc,
        23,
        "什么是 assistant-only loss？为什么需要它？",
        [
            "ChatML 样本通常包含 system、user 和 assistant 三部分。训练目标是让模型根据案件材料和指令生成目标文书，因此 system 和 user token 只作为上下文，不应该计算监督损失。",
            "实现上，对非 assistant 区间的 label 设为 ignore index，例如 -100，只在 assistant 输出 token 上计算交叉熵。这样可以避免模型把输入材料复述任务当成生成目标，也减少训练目标偏移。",
        ],
        formula="labels[token ∉ assistant_span] = -100\nloss = CrossEntropy(logits[assistant_span], labels[assistant_span])",
    )
    add_qa(
        doc,
        24,
        "为什么不用 BLEU、ROUGE 作为核心指标？",
        [
            "法律文书存在多种合理表述，文本相似度高不代表事实正确，文本相似度低也不一定质量差。更重要的是结构是否完整、主体和金额是否一致、是否产生无依据内容。",
            "因此我使用面向业务风险的指标：结构完整率检查必需段落和字段，主体一致率检查当事人角色及名称，幻觉率统计无案件依据的事实或结论。文本相似度可以作为辅助，但不应替代事实指标。",
        ],
        tips="幻觉率的分母是什么？由规则、模型 Judge 还是人工标注？",
        fill_item="三个指标的严格定义、评估脚本逻辑、人工复核方式和样本量。",
    )
    add_qa(
        doc,
        25,
        "如何保证 Base / SFT 对比公平？",
        [
            "必须固定验证集、prompt、system 指令、解码参数、最大生成长度和后处理逻辑。否则提升可能来自 prompt 或 decoding 变化，而不是微调。",
            "我会保存模型版本、推理配置和逐样本输出，做成可回归的对比报告。对于波动较大的生成任务，还可以在关键集上多次采样或使用确定性解码，减少随机性干扰。",
        ],
        fill_item="实际 temperature、top_p、max_new_tokens、验证集规模，以及是否使用 greedy decoding。",
    )
    add_qa(
        doc,
        26,
        "如何避免小样本 SFT 过拟合模板？",
        [
            "第一，做任务多样化，不让所有样本都只有完整文书生成。第二，按案件拆分数据，避免相似案件或同案派生任务跨集合。第三，监控验证集指标与训练 loss 的分离，结合 early stopping。第四，对主体、金额、拒答和信息缺失场景单独分桶评估。",
            "如果资源允许，还可以做 LoRA rank、epoch 和学习率的消融，检查模型是否出现固定句式复制、过度补全或拒答边界变差。",
        ],
    )
    add_qa(
        doc,
        27,
        "幻觉率从 0.26 降到 0.09，你认为原因是什么？",
        [
            "主要原因不是单一训练技巧，而是监督信号更贴近风险点。数据里加入了主体和金额纠错、信息不足拒答等任务，使模型学会在证据不足时不补全；assistant-only loss 也让训练目标更聚焦于正确输出。",
            "为了证明因果关系，最好做消融：只做完整文书 SFT、加入多任务拆解、再加入 assistant-only loss，比较各阶段指标变化。简历给出了最终结果，面试前我会准备实际消融记录；如果当时没有做过，就明确说明并给出补实验设计。",
        ],
        fill_item="是否有消融实验；如果没有，准备如实说明，并给出可执行的补实验方案。",
    )
    add_qa(
        doc,
        28,
        "如果文书生成系统进入生产，你会怎么做质量保障？",
        [
            "我会采用结构化输入、字段级校验和生成后审查。生成前校验主体、金额、请求和证据字段完整性；生成后检查必需段落、主体和金额一致性、关键结论是否有依据；高风险或低置信样本进入人工复核。",
            "此外要记录模型版本、prompt、输入摘要和审查结果，支持回溯。法律类场景中，系统更适合作为辅助起草工具，而不是无审核自动出具最终文书。",
        ],
    )


def add_foundations(doc):
    doc.add_heading("五、基础原理与场景题", level=1)
    add_qa(
        doc,
        29,
        "SFT、DPO、PPO、GRPO 分别适合解决什么问题？",
        [
            "SFT 用标注答案学习基本能力和格式，是建立初始策略的常用方式。DPO 直接利用偏好对优化相对偏好，不需要显式训练 reward model。PPO 使用 reward 和 critic 估计优势，适合在线或采样式策略优化，但工程复杂度更高。GRPO 使用同组样本的相对 reward 构造优势，不单独依赖 value model，适合对同一问题生成多条可评分轨迹的场景。",
            "实际项目中通常不是四选一，而是先 SFT 建立基础行为，再根据是否有偏好对、可验证 reward、在线 rollout 和算力条件选择后训练方法。",
        ],
    )
    add_qa(
        doc,
        30,
        "entropy bonus 和 KL penalty 有什么区别？",
        [
            "entropy bonus 鼓励当前策略保持探索，避免动作分布过早变得尖锐。KL penalty 约束新策略不要偏离参考策略过远，避免训练不稳定或能力漂移。",
            "两者作用不同：一个关注策略自身的不确定性，一个关注相对参考模型的偏移。在 Agent 任务中，可以同时监控动作熵、KL、任务成功率和高风险错误率，避免只优化单一指标。",
        ],
    )
    add_qa(
        doc,
        31,
        "如果 Agent 成功率低，你会如何系统排查？",
        [
            "我会先分层。第一层看环境和工具：动作解析、工具返回、超时和状态更新是否正确。第二层看数据：case 是否覆盖关键路径，reference path 与规则是否一致。第三层看策略：动作分布、平均步数、失败路径和分桶成功率。第四层看 reward：是否漏奖励、错惩罚或被利用。第五层看训练：KL、熵、梯度、组内 reward 方差和 rollout 多样性。",
            "排查顺序很重要。环境或 reward 有漏洞时，直接加训练步数通常会让错误策略学得更牢。",
        ],
    )
    add_qa(
        doc,
        32,
        "如何设计 Agent 的离线回归集？",
        [
            "回归集应按业务风险分层，而不是只做随机采样。至少包括正常办理、信息缺失、资格不满足、材料缺失、风险命中、工具异常、歧义输入和越权请求。每个 case 要有预期关键动作、允许路径、禁止动作和最终结果。",
            "评估时同时看结果正确性和过程合规性。对于高风险样本，错误 Submit 应单独统计并设为红线指标。",
        ],
    )
    add_qa(
        doc,
        33,
        "如果让你为电商售后 Agent 设计 reward，你会怎么迁移？",
        [
            "我会先把政务事项 schema 替换成售后 case schema：订单状态、商品类型、物流、退款规则、证据、用户诉求和风险标记。动作空间可以包括 Order_Query、Policy_Search、Evidence_Check、Refund_Calc、Risk_Check、Submit_Resolution 和 Escalate_Human。",
            "reward 仍然分层：规则 Verifier 校验订单事实、金额、权限和必要步骤；Judge 只评估回复是否清晰可执行；高风险动作例如错误退款、越权承诺必须有 hard gate。迁移的核心不是照搬动作，而是照搬“流程结构化 + 硬事实优先 + 轨迹可回放”的方法。",
        ],
    )
    add_qa(
        doc,
        34,
        "如果线上发现 Agent 调用工具次数过多，怎么优化？",
        [
            "先区分冗余调用和合理调用。可以统计每类 case 的必要工具集合、重复调用率、无效调用率、平均步数和工具耗时。对于重复查询，考虑缓存和状态记忆；对于无效调用，强化动作前置条件；对于路径过长，在不损害成功率的前提下加入轻量 step cost。",
            "step cost 不能过强，否则模型可能跳过必要工具。需要和成功率、漏调率、高风险错误率一起做 Pareto 权衡。",
        ],
    )
    add_qa(
        doc,
        35,
        "给你一份 trajectory log，你会写什么分析脚本？",
        [
            "我会按 case_type 和训练 step 分桶，计算成功率、final action 分布、每类工具调用率、漏调率、重复调用率、平均步数、组内 reward 均值与方差、路径去重率，以及高风险错误 Submit 数量。",
            "然后抽取典型失败轨迹：高 reward 但失败的样本用于找 reward hacking，低 reward 但正确的样本用于找错惩罚，动作高度同质的组用于找策略坍塌。输出最好能链接到完整 trace，便于回放。",
        ],
        formula="groupby(case_type, checkpoint) -> success_rate, action_histogram,\nrequired_tool_recall, duplicate_call_rate, reward_mean, reward_std,\nunique_path_ratio, wrong_submit_count",
    )


def add_behavior(doc):
    doc.add_heading("六、行为面试与反问", level=1)
    add_qa(
        doc,
        36,
        "讲一个你遇到的困难，以及你如何解决。",
        [
            "我会讲 reward hacking。最初模型在材料缺失场景中输出很完整的泛化建议，表面上 Judge 分数不错，但流程上漏掉了 Material_Check。这个问题不是继续调 prompt 能解决的。",
            "我回溯 trajectory log 和 Judge 理由，把问题拆成必要工具漏调、Verifier 覆盖不足和 Judge 权重过高三部分。然后加入 missing-tool penalty，约束 Submit 前的必要步骤，并降低 Judge 权重。这个经历让我认识到，应用算法优化要关注模型实际行为，而不是只看最终文本或总 reward。",
        ],
    )
    add_qa(
        doc,
        37,
        "如果需求不清晰，你如何和业务方合作？",
        [
            "我会先把模糊需求转换成可验收的问题：哪些结果绝对不能错，哪些步骤必须执行，哪些场景可以转人工，哪些指标代表体验。然后用少量典型 case 和失败样本对齐规则。",
            "对 Agent 项目尤其重要，因为“回复看起来不错”不能等同于任务完成。业务方提供风险边界，算法侧把它转成 schema、Verifier、回归集和日志指标，双方才能迭代。",
        ],
    )
    add_qa(
        doc,
        38,
        "你希望在实习中补齐什么能力？",
        [
            "我希望重点补齐三方面。第一，大规模训练和评估：更规范的实验管理、消融和回归。第二，线上系统能力：真实流量下的时延、成本、灰度和监控。第三，复杂 Agent 的工程治理：工具权限、异常处理、长程记忆和安全兜底。",
            "我已经有从业务流程到训练信号的实践，希望在更大规模场景中把这套方法做得更完整。",
        ],
    )
    doc.add_heading("建议向面试官提问", level=2)
    for text in (
        "团队当前的 Agent 项目更关注训练侧优化、评估体系，还是线上工具链和业务落地？",
        "对于实习生，入组后最常见的第一个问题是什么？更看重算法实验还是端到端交付？",
        "团队目前如何评估 Agent 的过程正确性？是否已经有 trajectory 回放、分桶回归和线上监控体系？",
        "岗位后续最希望补齐的能力是什么？",
    ):
        p = doc.add_paragraph(style="Answer")
        p.paragraph_format.left_indent = Inches(0.18)
        set_run_font(p.add_run(text), size=10.5, color=INK)


def add_checklist(doc, bullet_num_id):
    doc.add_heading("七、面试前必须补齐的真实数据", level=1)
    add_callout(
        doc,
        "重要提醒",
        "下面这些数据在简历中没有披露，但阿里算法面试中很可能被追问。请根据实验日志填写真实值；没有记录的项目就明确说明，并给出补实验方案。",
        fill=WARN,
    )
    rows = [
        ("Agent 数据", "case 数量、事项类型数、训练 / 验证拆分、难例占比", "待填写"),
        ("GRPO 配置", "底座模型、group size、rollout 数、学习率、KL、entropy 系数", "待填写"),
        ("Reward 配置", "各分项权重、hard gate、wrong-submit 与 missing-tool penalty", "待填写"),
        ("坍塌修复", "Refuse / Submit 分布、组内 reward 方差、必要工具调用率变化", "待填写"),
        ("Hacking 修复", "Material_Check 调用率、错误 Submit 率、Judge 权重变化", "待填写"),
        ("SFT 数据", "各任务类型比例、训练 / 验证 / 测试拆分、人工抽检比例", "待填写"),
        ("LoRA 配置", "Base / Instruct 版本、rank、alpha、dropout、lr、epoch、seq length", "待填写"),
        ("SFT 评估", "验证集规模、解码参数、指标脚本、人工复核方式", "待填写"),
        ("工程环境", "GPU、训练时长、推理框架、延迟、是否部署或仅离线验证", "待填写"),
    ]
    add_table(doc, ["模块", "需要准备的真实信息", "状态"], rows, [1600, 6200, 1560], header_fill=LIGHT_BLUE, font_size=9.2)

    doc.add_heading("最后 30 分钟复习顺序", level=2)
    for text in (
        "背熟 90 秒自我介绍，确保两段项目都有明确的个人贡献。",
        "补齐黄色提示中的真实参数和修复前后数据。",
        "用“现象 → 指标 → 定位 → 修复 → 验证”复述策略坍塌和 reward hacking。",
        "确认 assistant-only loss、GRPO 相对优势、entropy bonus 与 KL penalty 的区别。",
        "准备一个可迁移场景：电商售后、客服或企业服务 Agent 的 schema、动作和 reward。",
    ):
        add_bullet(doc, text, bullet_num_id)

    doc.add_heading("一页速记", level=2)
    add_table(
        doc,
        ["关键词", "一句话表达"],
        [
            ("Agent 环境", "case schema + 结构化动作 + 可回放 trajectory log"),
            ("GRPO", "同一 case 多 rollout，用组内相对 reward 构造优势"),
            ("坍塌", "动作集中导致路径差异和组内对比信号变弱"),
            ("Reward hacking", "模型利用 Judge 表达偏好，绕过必要 Material_Check"),
            ("修复原则", "硬事实优先，必要步骤 hard gate，表达质量只做补充"),
            ("SFT 数据", "按案件先拆分集合，再做受控多任务拆解，防止泄漏"),
            ("assistant-only loss", "只对 assistant 输出 token 计算监督损失"),
            ("法律文书评估", "结构、主体、金额与幻觉控制优先于文本相似度"),
        ],
        [1900, 7460],
        header_fill=LIGHT_BLUE,
        font_size=9.3,
        add_spacer=False,
    )
    add_terminal_marker(doc)


def set_document_metadata(doc):
    props = doc.core_properties
    props.title = "阿里算法岗日常实习｜简历定制面试题与参考回答"
    props.subject = "Agent 算法 / 应用算法面试准备"
    props.author = "Codex"
    props.keywords = "阿里, 算法实习, Agentic RL, GRPO, Qwen3-8B, LoRA SFT"
    props.comments = "基于候选人简历定制生成"


def configure_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)


def build():
    doc = Document()
    configure_page(doc)
    setup_styles(doc)
    set_document_metadata(doc)
    bullet_num_id = add_numbering(doc, bullet=True)
    decimal_num_id = add_numbering(doc, bullet=False)
    _ = decimal_num_id
    add_header_footer(doc)

    add_cover(doc)
    page_break(doc)
    add_overview(doc, bullet_num_id)
    page_break(doc)
    add_opening(doc)
    add_agent_rl(doc)
    page_break(doc)
    add_sft(doc)
    page_break(doc)
    add_foundations(doc)
    page_break(doc)
    add_behavior(doc)
    page_break(doc)
    add_checklist(doc, bullet_num_id)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
