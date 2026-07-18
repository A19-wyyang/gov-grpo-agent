from pathlib import Path
from zipfile import ZipFile
import sys

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


DOCX = Path(sys.argv[1]) if len(sys.argv) > 1 else (
    Path(__file__).resolve().parents[1]
    / "deliverables"
    / "阿里算法岗日常实习_简历定制面试题与参考回答_杨文宇.docx"
)

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def assert_equal(actual, expected, label):
    if actual != expected:
        raise AssertionError(f"{label}: expected {expected!r}, got {actual!r}")


def style_spacing(style):
    p_pr = style._element.find(qn("w:pPr"))
    spacing = p_pr.find(qn("w:spacing")) if p_pr is not None else None
    return {
        "before": spacing.get(qn("w:before")) if spacing is not None else None,
        "after": spacing.get(qn("w:after")) if spacing is not None else None,
        "line": spacing.get(qn("w:line")) if spacing is not None else None,
        "rule": spacing.get(qn("w:lineRule")) if spacing is not None else None,
    }


def audit_table_xml(root, label):
    tables = root.xpath(".//w:tbl", namespaces=NS)
    for idx, table in enumerate(tables, start=1):
        tbl_w = table.xpath("./w:tblPr/w:tblW/@w:w", namespaces=NS)
        tbl_ind = table.xpath("./w:tblPr/w:tblInd/@w:w", namespaces=NS)
        grid = [int(x) for x in table.xpath("./w:tblGrid/w:gridCol/@w:w", namespaces=NS)]
        if not tbl_w or not tbl_ind or not grid:
            raise AssertionError(f"{label} table {idx}: missing fixed geometry")
        if int(tbl_w[0]) != sum(grid):
            raise AssertionError(f"{label} table {idx}: tblW {tbl_w[0]} != grid sum {sum(grid)}")
        for row_idx, row in enumerate(table.xpath("./w:tr", namespaces=NS), start=1):
            cell_widths = [int(x) for x in row.xpath("./w:tc/w:tcPr/w:tcW/@w:w", namespaces=NS)]
            if cell_widths != grid:
                raise AssertionError(
                    f"{label} table {idx} row {row_idx}: tcW {cell_widths} != grid {grid}"
                )


def main():
    doc = Document(DOCX)
    section = doc.sections[0]

    assert_equal(section.page_width.twips, 12240, "page width")
    assert_equal(section.page_height.twips, 15840, "page height")
    assert_equal(section.top_margin.twips, 1440, "top margin")
    assert_equal(section.bottom_margin.twips, 1440, "bottom margin")
    assert_equal(section.left_margin.twips, 1440, "left margin")
    assert_equal(section.right_margin.twips, 1440, "right margin")
    assert_equal(section.header_distance.twips, 708, "header distance")
    assert_equal(section.footer_distance.twips, 708, "footer distance")

    expected_styles = {
        "Normal": {"after": "120", "line": "300"},
        "Heading 1": {"before": "360", "after": "200"},
        "Heading 2": {"before": "280", "after": "140"},
        "Heading 3": {"before": "200", "after": "100"},
    }
    for style_name, expected in expected_styles.items():
        spacing = style_spacing(doc.styles[style_name])
        for key, value in expected.items():
            assert_equal(spacing[key], value, f"{style_name} {key}")

    questions = [p for p in doc.paragraphs if p.style.name == "Question"]
    assert_equal(len(questions), 38, "question count")

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    required_headings = [
        "一、面试官视角：这份简历会怎样被追问",
        "三、政务办理 Agent：GRPO 策略优化",
        "四、仲裁文书生成：Qwen3-8B LoRA SFT",
        "七、面试前必须补齐的真实数据",
    ]
    for heading in required_headings:
        if heading not in headings:
            raise AssertionError(f"missing heading: {heading}")

    with ZipFile(DOCX) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        numbering_xml = etree.fromstring(archive.read("word/numbering.xml"))
        footer_xml = etree.fromstring(archive.read("word/footer1.xml"))

    numbered_paragraphs = document_xml.xpath(".//w:p[w:pPr/w:numPr]", namespaces=NS)
    if len(numbered_paragraphs) < 8:
        raise AssertionError(f"expected real numbered list paragraphs, found {len(numbered_paragraphs)}")
    bullet_levels = numbering_xml.xpath(".//w:lvl[w:numFmt/@w:val='bullet']", namespaces=NS)
    if not bullet_levels:
        raise AssertionError("missing custom bullet numbering definition")

    fake_bullets = [
        p.text for p in doc.paragraphs
        if p.text.strip().startswith(("•", "-", "●"))
    ]
    if fake_bullets:
        raise AssertionError(f"found fake bullets: {fake_bullets[:3]}")

    audit_table_xml(document_xml, "document")
    audit_table_xml(footer_xml, "footer")

    page_fields = footer_xml.xpath(".//w:instrText[contains(text(), 'PAGE')]", namespaces=NS)
    if not page_fields:
        raise AssertionError("footer PAGE field missing")

    table_text = [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    ]
    full_text = "\n".join([*(p.text for p in doc.paragraphs), *table_text])
    for forbidden in ("Project Lighthouse", "Riverbend", "TODO", "lorem ipsum"):
        if forbidden.lower() in full_text.lower():
            raise AssertionError(f"sample or placeholder text remains: {forbidden}")
    for required in ("待填写", "reward hacking", "assistant-only loss", "Material_Check"):
        if required.lower() not in full_text.lower():
            raise AssertionError(f"required interview content missing: {required}")

    print("OOXML structural audit passed")
    print(f"file={DOCX}")
    print(f"paragraphs={len(doc.paragraphs)} headings={len(headings)} questions={len(questions)}")
    print(f"tables={len(doc.tables)} numbered_paragraphs={len(numbered_paragraphs)} page_field=present")


if __name__ == "__main__":
    main()
